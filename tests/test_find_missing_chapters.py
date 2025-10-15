import sys
from pathlib import Path

from docx import Document

sys.path.append(str(Path(__file__).resolve().parents[1]))
from cod import find_missing_chapters  # noqa: E402


def _create_document(path, headings):
    document = Document()
    for heading in headings:
        document.add_paragraph(f"Глава {heading}")
        document.add_paragraph("Content")
    document.save(path)


def test_missing_simple_numbering(tmp_path):
    doc_path = tmp_path / "simple.docx"
    _create_document(doc_path, ["1", "2", "4"])

    issues = find_missing_chapters(str(doc_path))

    assert issues["missing"] == ["Глава 3"]
    assert issues["duplicates"] == []
    assert issues["unexpected"] == []


def test_missing_decimal_numbering(tmp_path):
    doc_path = tmp_path / "decimal.docx"
    _create_document(doc_path, ["1.1", "1.2", "2.1", "3.1"])

    issues = find_missing_chapters(str(doc_path))

    assert issues["missing"] == ["Глава 2.2", "Глава 3.2"]
    assert issues["duplicates"] == []
    assert issues["unexpected"] == []


def test_no_missing_chapters(tmp_path):
    doc_path = tmp_path / "complete.docx"
    _create_document(doc_path, ["1", "2", "3"])

    issues = find_missing_chapters(str(doc_path))

    assert issues == {"missing": [], "duplicates": [], "unexpected": []}


def test_duplicate_chapters(tmp_path):
    doc_path = tmp_path / "duplicates.docx"
    _create_document(doc_path, ["1", "2", "2", "3"])

    issues = find_missing_chapters(str(doc_path))

    assert issues["duplicates"] == ["Глава 2"]
    assert issues["missing"] == []
    assert issues["unexpected"] == []


def test_unexpected_minor(tmp_path):
    doc_path = tmp_path / "unexpected.docx"
    _create_document(
        doc_path,
        [
            "1.1",
            "1.2",
            "2.1",
            "2.2",
            "3.1",
            "3.2",
            "4.1",
            "4.2",
            "4.3",
        ],
    )

    issues = find_missing_chapters(str(doc_path))

    assert issues["unexpected"] == ["Глава 4.3"]
    assert issues["missing"] == []
    assert issues["duplicates"] == []

