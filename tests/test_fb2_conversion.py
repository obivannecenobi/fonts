from docx import Document

from fb2_converter import convert_docx_to_fb2


def test_convert_docx_to_fb2_creates_fb2_file(tmp_path):
    docx_path = tmp_path / "chapter.docx"
    document = Document()
    document.add_heading("Chapter 1")
    document.add_paragraph("Content line")
    document.save(docx_path)

    output_dir = tmp_path / "fb2"
    output_path = convert_docx_to_fb2(str(docx_path), str(output_dir))

    assert output_path.exists()
    content = output_path.read_text(encoding="utf-8")
    assert "<p>Chapter 1</p>" in content
    assert "<p>Content line</p>" in content

    # Validate FB2 structure
    from lxml import etree

    tree = etree.parse(str(output_path))
    root = tree.getroot()
    assert root.tag.endswith("FictionBook")
    paragraphs = root.findall(".//{http://www.gribuser.ru/xml/fictionbook/2.0}p")
    texts = [p.text for p in paragraphs]
    assert "Chapter 1" in texts
    assert "Content line" in texts
