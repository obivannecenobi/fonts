import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn

sys.path.append(str(Path(__file__).resolve().parents[1]))

from cod import collect_formatted_separators, fix_formatted_separator

VML_NS = "urn:schemas-microsoft-com:vml"
OFFICE_NS = "urn:schemas-microsoft-com:office:office"
W_NS = nsmap["w"]
W_PICT = f"{{{W_NS}}}pict"
W_DRAWING = f"{{{W_NS}}}drawing"


def _append_vml_horizontal_rule(paragraph):
    if "v" not in nsmap:
        nsmap["v"] = VML_NS
    if "o" not in nsmap:
        nsmap["o"] = OFFICE_NS

    run = paragraph.add_run()
    pict = OxmlElement("w:pict")
    shape = OxmlElement("v:rect")
    shape.set(qn("o:hr"), "t")
    pict.append(shape)
    run._r.append(pict)


def _add_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def test_collect_formatted_separators_detects_vml_shape():
    document = Document()
    paragraph = document.add_paragraph()
    _append_vml_horizontal_rule(paragraph)

    separators = collect_formatted_separators(document)

    assert len(separators) == 1
    index, found_paragraph = separators[0]
    assert index == 1
    assert found_paragraph._p is paragraph._p


def test_fix_formatted_separator_replaces_shape_and_border():
    document = Document()
    paragraph = document.add_paragraph()
    _append_vml_horizontal_rule(paragraph)
    _add_paragraph_border(paragraph)

    fix_formatted_separator(paragraph)

    assert paragraph.text == "***\u200B"

    assert not any(child.tag == W_PICT for child in paragraph._p.iter())
    assert not any(child.tag == W_DRAWING for child in paragraph._p.iter())

    p_pr = paragraph._p.pPr
    if p_pr is not None:
        assert p_pr.find(qn("w:pBdr")) is None
