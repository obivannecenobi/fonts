import os
import sys
from pathlib import Path

from docx import Document

sys.path.append(str(Path(__file__).resolve().parents[1]))
from cod import check_english_words


def test_check_english_words(tmp_path):
    path = tmp_path / "english.docx"
    doc = Document()
    doc.add_paragraph("Привет Hello мир")
    doc.add_paragraph("Это Test документ")
    doc.add_paragraph("Другие слова: World, Hello")
    doc.save(path)

    words = check_english_words(str(path))
    assert words == {
        "Hello": [(1, 8), (3, 22)],
        "Test": [(2, 5)],
        "World": [(3, 15)],
    }

    save_path = tmp_path / "words.txt"
    lines = [
        f"{word} - {', '.join(str(p) for p, _ in positions)}"
        for word, positions in words.items()
    ]
    with open(save_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    assert (
        save_path.read_text(encoding="utf-8")
        == "Hello - 1, 3\nTest - 2\nWorld - 3"
    )

