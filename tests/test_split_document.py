import os
import re
import sys
from pathlib import Path
from unittest.mock import patch

from docx import Document

sys.path.append(str(Path(__file__).resolve().parents[1]))
from cod import split_document, split_chapters_into_two


def create_sample_doc(path):
    doc = Document()
    doc.add_paragraph("Глава 1")
    doc.add_paragraph("Text for chapter 1")
    doc.add_paragraph("Глава 1.1")
    doc.add_paragraph("Text for chapter 1.1")
    doc.add_paragraph("Глава 2")
    doc.add_paragraph("Text for chapter 2")
    doc.save(path)


def test_split_document(tmp_path):
    source = tmp_path / "source.docx"
    create_sample_doc(source)

    with patch("cod.filedialog.askdirectory", return_value=str(tmp_path)):
        created_files = split_document(str(source))

    expected_names = {"Глава 1.docx", "Глава 1.1.docx", "Глава 2.docx"}
    assert {os.path.basename(p) for p in created_files} == expected_names

    heading_pattern = re.compile(r"^Глава\s+\d+(?:\.\d+)?")
    for file_name in expected_names:
        doc = Document(tmp_path / file_name)
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        assert all(not heading_pattern.match(t) for t in texts)


def create_duplicate_doc(path):
    doc = Document()
    doc.add_paragraph("Глава 1")
    doc.add_paragraph("Text for chapter 1")
    doc.add_paragraph("Глава 1")
    doc.add_paragraph("Another text")
    doc.save(path)


def test_split_document_duplicate_titles(tmp_path):
    source = tmp_path / "dup_source.docx"
    create_duplicate_doc(source)

    with patch("cod.filedialog.askdirectory", return_value=str(tmp_path)):
        created_files = split_document(str(source))

    expected_names = {"Глава 1.docx", "Глава 1 (2).docx"}
    assert {os.path.basename(p) for p in created_files} == expected_names


def create_even_split_doc(path):
    doc = Document()
    doc.add_paragraph("Глава 1")
    doc.add_paragraph("Первый абзац первой главы")
    doc.add_paragraph("Второй абзац первой главы")
    doc.add_paragraph("Глава 2")
    doc.add_paragraph("Первый абзац второй главы")
    doc.add_paragraph("Второй абзац второй главы")
    doc.save(path)


def test_split_chapters_into_two_creates_pair_documents(tmp_path):
    source = tmp_path / "even.docx"
    create_even_split_doc(source)

    created, skipped = split_chapters_into_two(str(source), str(tmp_path))

    expected_names = {
        "Глава 1.1.docx",
        "Глава 1.2.docx",
        "Глава 2.1.docx",
        "Глава 2.2.docx",
    }

    assert {os.path.basename(p) for p in created} == expected_names
    assert skipped == []

    contents = {
        Path(path).stem: [p.text for p in Document(path).paragraphs if p.text]
        for path in created
    }

    assert contents["Глава 1.1"] == ["Первый абзац первой главы"]
    assert contents["Глава 1.2"] == ["Второй абзац первой главы"]
    assert contents["Глава 2.1"] == ["Первый абзац второй главы"]
    assert contents["Глава 2.2"] == ["Второй абзац второй главы"]


def test_split_chapters_into_two_skips_short_chapter(tmp_path):
    source = tmp_path / "short.docx"
    doc = Document()
    doc.add_paragraph("Глава 1")
    doc.add_paragraph("Один абзац")
    doc.save(source)

    created, skipped = split_chapters_into_two(str(source), str(tmp_path))

    assert created == []
    assert skipped == ["1"]

