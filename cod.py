"""GUI generator for creating DOCX chapter files.

The application registers the bundled 'Cattedrale' font from the local fonts
directory so the font file must be available on disk.
"""

import ctypes
import json
import os
import sys
import time
import re
import tkinter as tk
from tkinter import filedialog, ttk
from tkinter import font as tkfont

import copy
import customtkinter as ctk
from docx import Document
from docx.oxml.ns import nsmap, qn
from docx.text.paragraph import Paragraph
from importlib import import_module, util
from typing import Any, Dict, List, Optional, Tuple

from rulate_uploader import upload_chapters


def split_document(file_path: str) -> List[str]:
    """Split a DOCX document into chapters based on heading pattern.

    Parameters
    ----------
    file_path: str
        Path to the source DOCX file.

    Returns
    -------
    List[str]
        List of paths to the created chapter files.
    """

    heading_pattern = re.compile(r"^Глава\s+\d+(?:\.\d+)?")
    output_dir = filedialog.askdirectory(
        title="Выберите выходной каталог",
        initialdir=os.path.dirname(file_path),
    )
    if not output_dir:
        output_dir = os.path.dirname(file_path)
    document = Document(file_path)
    current_doc = None
    current_title = ""
    created_files: List[str] = []

    def _unique_path(directory: str, filename: str) -> str:
        base, ext = os.path.splitext(filename)
        candidate = os.path.join(directory, filename)
        counter = 2
        while os.path.exists(candidate):
            candidate = os.path.join(directory, f"{base} ({counter}){ext}")
            counter += 1
        return candidate

    for para in document.paragraphs:
        text = para.text.strip()
        if heading_pattern.match(text):
            if current_doc is not None:
                sanitized = re.sub(r"[^\w\s.-]", "", current_title).strip()
                if not sanitized:
                    sanitized = "section"
                file_name = f"{sanitized}.docx"
                out_path = _unique_path(output_dir, file_name)
                current_doc.save(out_path)
                created_files.append(out_path)
            current_doc = Document()
            current_title = text
        else:
            if current_doc is not None:
                new_element = copy.deepcopy(para._element)
                current_doc._element.body.append(new_element)

    if current_doc is not None:
        sanitized = re.sub(r"[^\w\s.-]", "", current_title).strip()
        if not sanitized:
            sanitized = "section"
        file_name = f"{sanitized}.docx"
        out_path = _unique_path(output_dir, file_name)
        current_doc.save(out_path)
        created_files.append(out_path)

    return created_files


def split_chapters_into_two(file_path: str, output_dir: str) -> Tuple[List[str], List[str]]:
    """Split each detected chapter into two separate documents."""

    heading_pattern = re.compile(r"^Глава\s+(\d+)(?:\.(\d+))?", re.IGNORECASE)
    document = Document(file_path)
    created_files: List[str] = []
    skipped_chapters: List[str] = []
    current_label: Optional[str] = None
    current_paragraphs: List[Paragraph] = []

    def _unique_path(directory: str, filename: str) -> str:
        base, ext = os.path.splitext(filename)
        candidate = os.path.join(directory, filename)
        counter = 2
        while os.path.exists(candidate):
            candidate = os.path.join(directory, f"{base} ({counter}){ext}")
            counter += 1
        return candidate

    def _clear_document(doc: Document) -> None:
        while doc.paragraphs:
            element = doc.paragraphs[0]._element
            element.getparent().remove(element)

    def _save_split(label: str, paragraphs: List[Paragraph]) -> None:
        if len(paragraphs) < 2:
            skipped_chapters.append(label)
            return

        weights = [max(len(p.text.strip()), 1) for p in paragraphs]
        total_weight = sum(weights)
        target = total_weight / 2
        cumulative = 0
        split_index = len(paragraphs) // 2

        for index, weight in enumerate(weights, start=1):
            cumulative += weight
            if cumulative >= target:
                split_index = index
                break

        if split_index >= len(paragraphs):
            split_index = len(paragraphs) - 1
        if split_index <= 0:
            split_index = 1

        first_part = paragraphs[:split_index]
        second_part = paragraphs[split_index:]

        if not second_part:
            skipped_chapters.append(label)
            return

        for part_index, part in enumerate((first_part, second_part), start=1):
            new_doc = Document()
            _clear_document(new_doc)
            for paragraph in part:
                new_doc._element.body.append(copy.deepcopy(paragraph._element))

            chapter_name = f"Глава {label}.{part_index}"
            sanitized = re.sub(r"[^\w\s.-]", "", chapter_name).strip()
            if not sanitized:
                sanitized = f"chapter_{part_index}"
            filename = f"{sanitized}.docx"
            output_path = _unique_path(output_dir, filename)
            new_doc.save(output_path)
            created_files.append(output_path)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = heading_pattern.match(text)
        if match:
            if current_label is not None:
                _save_split(current_label, current_paragraphs)
            major = match.group(1)
            minor = match.group(2)
            label = major if minor is None else f"{major}.{minor}"
            current_label = label
            current_paragraphs = []
        elif current_label is not None:
            current_paragraphs.append(paragraph)

    if current_label is not None:
        _save_split(current_label, current_paragraphs)

    return created_files, skipped_chapters


def check_english_words(file_path: str) -> Dict[str, List[Tuple[int, int]]]:
    """Return mapping of English or mixed-language words to their positions."""

    document = Document(file_path)
    results: Dict[str, List[Tuple[int, int]]] = {}
    english_pattern = re.compile(r"[A-Za-z]+")
    mixed_pattern = re.compile(r"[A-Za-zА-Яа-яЁё]+")

    for p_idx, para in enumerate(document.paragraphs, start=1):
        text = para.text

        for match in english_pattern.finditer(text):
            word = match.group()
            char_pos = match.start() + 1  # 1-indexed character position
            results.setdefault(word, []).append((p_idx, char_pos))

        for match in mixed_pattern.finditer(text):
            word = match.group()
            if re.search(r"[A-Za-z]", word) and re.search(r"[А-Яа-яЁё]", word):
                char_pos = match.start() + 1
                results.setdefault(word, []).append((p_idx, char_pos))

    return {word: positions for word, positions in sorted(results.items())}


def find_duplicate_chapters(file_path: str) -> List[Tuple[List[str], str]]:
    """Find duplicated chapter contents in a DOCX document."""

    document = Document(file_path)
    heading_pattern = re.compile(r"^Глава\s+\d+(?:\.\d+)?", re.IGNORECASE)
    chapters: List[Tuple[str, str]] = []
    current_title = None
    current_content: List[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if heading_pattern.match(text):
            if current_title is not None:
                content_text = "\n".join(current_content).strip()
                chapters.append((current_title, content_text))
            current_title = text
            current_content = []
        elif current_title is not None:
            current_content.append(text)

    if current_title is not None:
        content_text = "\n".join(current_content).strip()
        chapters.append((current_title, content_text))

    content_map: Dict[str, List[str]] = {}
    for title, content in chapters:
        content_map.setdefault(content, []).append(title)

    duplicates: List[Tuple[List[str], str]] = []
    for content, titles in content_map.items():
        if len(titles) > 1:
            duplicates.append((titles, content))

    return duplicates


def _format_chapter_number(parts: Tuple[int, ...]) -> str:
    """Return chapter label from tuple representation."""

    if len(parts) == 1:
        return f"Глава {parts[0]}"
    return f"Глава {parts[0]}.{parts[1]}"


def find_missing_chapters(file_path: str) -> List[str]:
    """Return the list of missing chapter headings in a DOCX document."""

    document = Document(file_path)
    heading_pattern = re.compile(r"^Глава\s+(\d+(?:\.\d+)?)", re.IGNORECASE)
    raw_numbers: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = heading_pattern.match(text)
        if match:
            raw_numbers.append(match.group(1))

    if not raw_numbers:
        return []

    contains_decimal = any("." in value for value in raw_numbers)

    if not contains_decimal:
        numbers = [int(value) for value in raw_numbers]
        missing: List[str] = []
        expected = numbers[0]

        for current in numbers:
            while expected < current:
                missing.append(_format_chapter_number((expected,)))
                expected += 1
            expected = current + 1

        return missing

    chapters: List[Tuple[int, int]] = []
    for value in raw_numbers:
        parts = value.split(".")
        if len(parts) != 2:
            continue
        chapters.append((int(parts[0]), int(parts[1])))

    if not chapters:
        return []

    missing: List[str] = []
    major_order: List[int] = []
    majors: Dict[int, List[int]] = {}

    for major, minor in chapters:
        if major not in majors:
            majors[major] = []
            major_order.append(major)
        majors[major].append(minor)

    prev_major: int | None = None
    prev_max_minor: int | None = None

    for major in major_order:
        if prev_major is not None and major - prev_major > 1:
            limit = prev_max_minor if prev_max_minor is not None else 1
            limit = max(limit, 1)
            for gap_major in range(prev_major + 1, major):
                for minor in range(1, limit + 1):
                    missing.append(_format_chapter_number((gap_major, minor)))

        ordered_minors = list(dict.fromkeys(majors[major]))
        ordered_minors.sort()

        expected_minor = 1
        for minor in ordered_minors:
            while expected_minor < minor:
                missing.append(_format_chapter_number((major, expected_minor)))
                expected_minor += 1
            expected_minor = minor + 1

        current_max_minor = ordered_minors[-1] if ordered_minors else 0
        original_max_minor = current_max_minor

        if prev_max_minor is not None and current_max_minor < prev_max_minor:
            for value in range(current_max_minor + 1, prev_max_minor + 1):
                missing.append(_format_chapter_number((major, value)))

        prev_major = major
        if original_max_minor:
            prev_max_minor = original_max_minor
        elif prev_max_minor is None:
            prev_max_minor = 1

    return missing


_W_NS = nsmap["w"]
_WP_NS = nsmap.get("wp", "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing")
_PIC_NS = nsmap.get("pic", "http://schemas.openxmlformats.org/drawingml/2006/picture")
_VML_NS = "urn:schemas-microsoft-com:vml"
_OFFICE_NS = "urn:schemas-microsoft-com:office:office"

_W_PICT = f"{{{_W_NS}}}pict"
_W_DRAWING = f"{{{_W_NS}}}drawing"
_WP_DOCPR = f"{{{_WP_NS}}}docPr"
_PIC_CNVPR = f"{{{_PIC_NS}}}cNvPr"
_VML_SHAPE_TAGS = {
    f"{{{_VML_NS}}}shape",
    f"{{{_VML_NS}}}rect",
    f"{{{_VML_NS}}}oval",
    f"{{{_VML_NS}}}line",
}
_OFFICE_HR_ATTRIBUTES = {
    f"{{{_OFFICE_NS}}}hr",
    f"{{{_OFFICE_NS}}}hrstd",
    f"{{{_OFFICE_NS}}}hralign",
    f"{{{_OFFICE_NS}}}hrpct",
}
_OFFICE_HR_TAG = f"{{{_OFFICE_NS}}}hr"

_HORIZONTAL_RULE_KEYWORDS = (
    "horizontal line",
    "horizontal rule",
    "горизонтальная линия",
    "разделитель",
)


def _paragraph_has_horizontal_border(paragraph: Paragraph) -> bool:
    """Return True if the paragraph is rendered as a horizontal separator line."""

    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    border = p_pr.find(qn("w:pBdr"))
    if border is None:
        return False
    separator_tags = {
        qn("w:top"),
        qn("w:bottom"),
        qn("w:between"),
        qn("w:bar"),
    }
    return any(child.tag in separator_tags for child in border)


def _element_defines_horizontal_rule(element) -> bool:
    """Return True if the provided XML element corresponds to an HR shape."""

    for child in element.iter():
        if child.tag == _OFFICE_HR_TAG:
            return True

        if child.tag in _VML_SHAPE_TAGS:
            for attr in _OFFICE_HR_ATTRIBUTES:
                value = child.get(attr)
                if value and value.lower() not in {"f", "false", "0"}:
                    return True
            style = child.get("style")
            if style and "hr" in style.replace(" ", "").lower():
                return True
    return False


def _paragraph_contains_horizontal_rule_shape(paragraph: Paragraph) -> bool:
    """Return True if the paragraph contains a shape representing an HR line."""

    if paragraph.text.strip():
        return False

    for pict in paragraph._p.iter(_W_PICT):
        if _element_defines_horizontal_rule(pict):
            return True

    for drawing in paragraph._p.iter(_W_DRAWING):
        for doc_pr in drawing.iter(_WP_DOCPR):
            text = f"{doc_pr.get('title', '')} {doc_pr.get('descr', '')}".lower()
            if any(keyword in text for keyword in _HORIZONTAL_RULE_KEYWORDS):
                return True

        for cnv_pr in drawing.iter(_PIC_CNVPR):
            text = f"{cnv_pr.get('title', '')} {cnv_pr.get('descr', '')}".lower()
            if any(keyword in text for keyword in _HORIZONTAL_RULE_KEYWORDS):
                return True

        if _element_defines_horizontal_rule(drawing):
            return True

    return False


def collect_formatted_separators(document: Document) -> List[Tuple[int, Paragraph]]:
    """Return paragraphs representing auto-formatted separator lines."""

    results: List[Tuple[int, Paragraph]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        if _paragraph_has_horizontal_border(paragraph) or _paragraph_contains_horizontal_rule_shape(
            paragraph
        ):
            results.append((index, paragraph))
    return results


def fix_formatted_separator(paragraph: Paragraph) -> None:
    """Replace a formatted separator with plain text '***\u200B'."""

    p_element = paragraph._p
    for child in list(p_element):
        if child.tag != qn("w:pPr"):
            p_element.remove(child)

    paragraph.text = "***\u200B"
    p_pr = paragraph._p.get_or_add_pPr()
    border = p_pr.find(qn("w:pBdr"))
    if border is not None:
        p_pr.remove(border)

# Path to store window geometry
CONFIG_PATH = os.path.join(os.path.dirname(__file__), "window.cfg")

class CustomInputDialog(ctk.CTkToplevel):
    """Simple dialog asking the user for a single line of text."""

    def __init__(self, master, question: str, font: ctk.CTkFont, icon_path: str):
        super().__init__(master)
        self.icon_path = icon_path
        self.iconbitmap(self.icon_path)
        self.title("")
        self.resizable(False, False)
        self.configure(fg_color="#2f2f2f")
        self.result = None

        self._label = ctk.CTkLabel(
            self, text=question, text_color="#eeeeee", font=font
        )
        self._label.pack(padx=20, pady=(20, 10))

        self._entry = ctk.CTkEntry(
            self,
            fg_color="#ffffff",
            border_color="#2f2f2f",
            text_color="#303030",
            corner_radius=getattr(master, "entry_corner_radius", 12),
            border_width=0,
            font=font,
        )
        self._entry.pack(padx=20, pady=(0, 20))

        entry_height = getattr(master, "entry_height", None)
        if entry_height:
            self._entry.configure(height=entry_height)

        button_frame = ctk.CTkFrame(self, fg_color="#2f2f2f")
        button_frame.pack(padx=20, pady=(0, 20))

        button_text_color = getattr(master, "button_text_color", "#eeeeee")

        self._ok_button = ctk.CTkButton(
            button_frame,
            text="OK",
            command=self._ok,
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=button_text_color,
            corner_radius=getattr(master, "button_corner_radius", 12),
            border_width=0,
            font=font,
        )
        self._ok_button.pack(side="left", padx=(0, 10))

        self._cancel_button = ctk.CTkButton(
            button_frame,
            text="Cancel",
            command=self._cancel,
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=button_text_color,
            corner_radius=getattr(master, "button_corner_radius", 12),
            border_width=0,
            font=font,
        )
        self._cancel_button.pack(side="left")

        button_height = getattr(master, "button_height", None)
        if button_height:
            self._ok_button.configure(height=button_height)
            self._cancel_button.configure(height=button_height)

        hover_helper = getattr(master, "_apply_button_hover_effect", None)
        if callable(hover_helper):
            hover_helper(self._ok_button)
            hover_helper(self._cancel_button)

        self._entry.bind("<Return>", lambda event: self._ok())
        self.protocol("WM_DELETE_WINDOW", self._cancel)

        self.update_idletasks()
        master.update_idletasks()
        x = master.winfo_rootx() + (master.winfo_width() // 2) - (
            self.winfo_width() // 2
        )
        y = master.winfo_rooty() + (master.winfo_height() // 2) - (
            self.winfo_height() // 2
        )
        self.geometry(f"+{x}+{y}")

    def _ok(self) -> None:
        self.result = self._entry.get()
        self.destroy()

    def _cancel(self) -> None:
        self.result = None
        self.destroy()

    def get_input(self):
        self._entry.focus()
        self.grab_set()
        self.wait_window()
        return self.result


class Application(tk.Tk):
    def __init__(self):
        super().__init__()
        self.icon_path = os.path.join(
            os.path.dirname(__file__),
            "ChatGPT Image 15 авг. 2025 г., 20_42_09.ico",
        )
        self.iconbitmap(self.icon_path)

        self.config_data = self.load_config()

        # Путь к вашему шрифту
        font_path = os.path.join(
            os.path.dirname(__file__),
            "fonts",
            "Cattedrale[RUSbypenka220]-Regular.ttf",
        )

        # Регистрация и настройка шрифта
        if sys.platform.startswith("win"):
            ctypes.windll.gdi32.AddFontResourceExW(font_path, 0x10, 0)
            font_family = "Cattedrale [RUS by penka220]"
        else:
            # На системах Unix пытаемся зарегистрировать шрифт через tkfont.
            # Если это невозможно, используем системный шрифт по умолчанию.
            try:
                self._loaded_font = tkfont.Font(file=font_path)
                font_family = self._loaded_font.actual("family")
            except tk.TclError:
                font_family = tkfont.nametofont("TkDefaultFont").actual("family")
        default_font = tkfont.nametofont("TkDefaultFont")
        base_size = default_font.cget("size") + 6
        font_size = max(int(self.config_data.get("font_size", base_size)), base_size)
        custom_font = ctk.CTkFont(family=font_family, size=font_size)
        self.custom_font = custom_font
        self.config_data["font_size"] = str(font_size)
        self.option_add("*Font", custom_font)

        self.button_height = max(int(font_size * 2.1), 40)
        self.entry_height = self.button_height
        self.button_text_color = "#f2f2f2"
        self.neon_text_color = "#ffffff"
        self.button_fg_color = "#313131"
        self.button_hover_color = "#181818"
        self.button_border_color = "#ffffff"
        self.button_border_width = 2
        self.button_corner_radius = 20
        self.entry_corner_radius = 20

        self.style = ttk.Style(self)
        self.style.theme_use("clam")

        self.style.configure("Custom.TFrame", background="#2f2f2f")
        self.style.configure("Custom.TLabel", background="#2f2f2f", foreground="#eeeeee")

        # Создаем окно перед настройкой шрифта
        self.title("НЕЙРО-СТРАЖ")
        # Set window geometry
        min_width, min_height = 560, 640
        saved_geom = self.config_data.get("geometry", "")
        if saved_geom:
            width, height, position = self._parse_geometry(saved_geom)
            width = max(width, min_width)
            height = max(height, min_height)
            geometry_value = f"{width}x{height}{position}"
        else:
            geometry_value = f"{min_width}x{min_height}"
        self.geometry(geometry_value)
        self.minsize(min_width, min_height)
        self.configure(bg="#2f2f2f")  # Темно-серый фон
        self.resizable(True, True)
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

        # Создаем рамку для текста
        self.frame = ctk.CTkFrame(self, fg_color="#2f2f2f")
        self.frame.pack(padx=20, pady=20, expand=True, fill="both")

        # Заголовок
        header_font = ctk.CTkFont(
            family=custom_font.actual("family"),
            size=max(font_size + 10, 28),
            weight="bold",
        )
        header_frame = ctk.CTkFrame(self.frame, fg_color="#2f2f2f")
        header_frame.pack(fill="x", pady=(0, 10))

        self.label = ttk.Label(
            header_frame,
            text="НЕЙРО-СТРАЖ",
            font=header_font,
            style="Custom.TLabel",
        )
        self.label.pack(pady=20)

        self.groups_container = ctk.CTkFrame(self.frame, fg_color="#2f2f2f")
        self.groups_container.pack(expand=True, fill="both")

        self.generator_group = ctk.CTkFrame(self.groups_container, fg_color="#2f2f2f")
        self.generator_group.pack(fill="x", padx=10, pady=(0, 10))

        self.browse_button = ctk.CTkButton(
            self.generator_group,
            text="Выбрать папку",
            command=self.browse_folder,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self.browse_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.browse_button)

        self.path_entry = ctk.CTkEntry(
            self.generator_group,
            placeholder_text="Куда сейвим?",
            corner_radius=self.entry_corner_radius,
            fg_color="#ffffff",
            text_color="#303030",
            border_color="#2f2f2f",
            border_width=0,
            bg_color="#2f2f2f",
            font=self.custom_font,
            height=self.entry_height,
        )
        self.path_entry.pack(fill=tk.X, pady=(0, 8))
        self._create_generate_button()

        self._add_separator(self.groups_container)

        fix_group = ctk.CTkFrame(self.groups_container, fg_color="#2f2f2f")
        fix_group.pack(fill="x", padx=10, pady=(12, 10))

        self.duplicates_button = ctk.CTkButton(
            fix_group,
            text="Повторение",
            command=self.check_duplicate_chapters,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self.duplicates_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.duplicates_button)

        self.separator_button = ctk.CTkButton(
            fix_group,
            text="Разделители",
            command=self.find_formatted_separators,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self.separator_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.separator_button)

        self.numbering_button = ctk.CTkButton(
            fix_group,
            text="Нумерация",
            command=self.check_chapter_numbering,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self.numbering_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.numbering_button)

        self.artifacts_button = ctk.CTkButton(
            fix_group,
            text="Артефакты",
            command=self.check_english_words,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self.artifacts_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.artifacts_button)

        self.split_even_button = ctk.CTkButton(
            fix_group,
            text="Разделить",
            command=self.split_chapters_evenly,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self.split_even_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.split_even_button)

        self.split_button = ctk.CTkButton(
            fix_group,
            text="Разбить",
            command=self.split_document,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self.split_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.split_button)

        self._add_separator(self.groups_container)

        convert_group = ctk.CTkFrame(self.groups_container, fg_color="#2f2f2f")
        convert_group.pack(fill="x", padx=10, pady=(12, 10))

        self.convert_button = ctk.CTkButton(
            convert_group,
            text="Законвертить",
            command=self.convert_docx_to_fb2,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self.convert_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.convert_button)

        # Button to upload chapters to Rulate (hidden by default)
        self.upload_button = ctk.CTkButton(
            self.groups_container,
            text="Залить на Rulate",
            command=self.open_upload_dialog,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        self._apply_button_hover_effect(self.upload_button)

        self.update_idletasks()
        width, height, position = self._parse_geometry(self.geometry())
        content_height = self.frame.winfo_reqheight() + 40
        max_height = self.winfo_screenheight() - 80
        adjusted_min_height = max(min_height, min(content_height, max_height))
        self.minsize(min_width, adjusted_min_height)
        target_height = min(max(height, content_height), max_height)
        if target_height > height:
            self.geometry(f"{width}x{target_height}{position}")

    def _apply_button_hover_effect(self, button: ctk.CTkButton) -> None:
        button.configure(
            fg_color=self.button_fg_color,
            hover_color=self.button_hover_color,
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            border_width=0,
            border_color=self.button_border_color,
        )

        default_color = self.button_text_color

        def _on_enter(_: tk.Event) -> None:  # type: ignore[override]
            button.configure(
                text_color=self.neon_text_color,
                border_width=self.button_border_width,
            )

        def _on_leave(_: tk.Event) -> None:  # type: ignore[override]
            button.configure(
                text_color=default_color,
                border_width=0,
            )

        button.bind("<Enter>", _on_enter)
        button.bind("<Leave>", _on_leave)

    def _add_separator(self, parent: tk.Widget) -> None:
        container = tk.Frame(parent, bg="#2f2f2f")
        container.pack(fill="x")

        line = tk.Frame(container, bg="#ffffff", height=1)
        line.pack(pady=(12, 12))

        def _resize_line(event: tk.Event) -> None:  # type: ignore[override]
            width = int(event.width * (2 / 3))
            line.configure(width=max(width, 1))

        container.bind("<Configure>", _resize_line)

    def _create_generate_button(self) -> None:
        if hasattr(self, "ask_button"):
            self.ask_button.destroy()

        width = max(int(self.button_height * 5), 220)
        self.ask_button = ctk.CTkButton(
            self.generator_group,
            text="Сгенерировать",
            command=self.ask_questions,
            width=width,
            height=self.button_height,
            corner_radius=self.button_corner_radius,
            fg_color=self.button_fg_color,
            hover_color=self.button_hover_color,
            bg_color="#2f2f2f",
            text_color=self.button_text_color,
            font=self.custom_font,
        )
        self.ask_button.pack(pady=(0, 8))
        self._apply_button_hover_effect(self.ask_button)

    def _parse_geometry(self, geometry: str) -> tuple[int, int, str]:
        size_part = geometry
        position_part = ""
        for index, char in enumerate(geometry):
            if char in "+-":
                size_part = geometry[:index]
                position_part = geometry[index:]
                break

        try:
            width_str, height_str = size_part.split("x", 1)
            width = int(width_str)
            height = int(height_str)
        except ValueError:
            width, height = 0, 0

        return width, height, position_part


    def browse_folder(self):
        folder_selected = filedialog.askdirectory(title="Выберите папку для сохранения")
        if folder_selected:
            self.path_entry.delete(0, tk.END)
            self.path_entry.insert(0, folder_selected)

    def split_document(self):
        file_path = filedialog.askopenfilename(
            title="Выберите документ",
            filetypes=[("Word Documents", "*.docx")],
        )
        if not file_path:
            return

        created = split_document(file_path)
        if created:
            self.show_message(f"Создано {len(created)} файлов")

    def split_chapters_evenly(self):
        file_path = filedialog.askopenfilename(
            title="Выберите документ",
            filetypes=[("Word Documents", "*.docx")],
        )
        if not file_path:
            return

        output_dir = filedialog.askdirectory(
            title="Выберите папку для сохранения",
            initialdir=os.path.dirname(file_path),
        )
        if not output_dir:
            self.show_error("Папка для сохранения не выбрана.")
            return

        created, skipped = split_chapters_into_two(file_path, output_dir)

        if not created:
            message = "Не удалось разделить главы."
            if skipped:
                message += "\n" + "\n".join(
                    f"Глава {label}" for label in skipped
                )
            self.show_error(message)
            return

        message_lines = [f"Создано {len(created)} файлов"]
        if skipped:
            skipped_text = ", ".join(f"Глава {label}" for label in skipped)
            message_lines.append(f"Не удалось разделить: {skipped_text}")

        self.show_message("\n".join(message_lines))

    def convert_docx_to_fb2(self):
        if util.find_spec("fb2_converter") is None:
            self.show_error(
                "Модуль fb2_converter не найден. Убедитесь, что файлы проекта на месте."
            )
            return
        if util.find_spec("lxml") is None:
            self.show_error(
                "Библиотека lxml не установлена. Установите зависимости из requirements.txt."
            )
            return

        converter_module = import_module("fb2_converter")

        files = filedialog.askopenfilenames(
            title="Выберите документы для конвертации",
            filetypes=[("Word Documents", "*.docx")],
        )
        if not files:
            return

        destination = filedialog.askdirectory(
            title="Выберите папку для FB2 файлов",
            initialdir=os.path.dirname(files[0]),
        )
        if not destination:
            self.show_error("Папка для сохранения не выбрана.")
            return

        success_count = 0
        failures: List[Tuple[str, str]] = []

        convert_docx_to_fb2_file = converter_module.convert_docx_to_fb2

        for docx_path in files:
            try:
                convert_docx_to_fb2_file(docx_path, destination)
                success_count += 1
            except Exception as exc:  # noqa: BLE001 - surface errors to the user
                failures.append((docx_path, str(exc)))

        if success_count:
            self.show_popup(
                f"Сконвертировано {success_count} файл(ов) в {destination}"
            )

        if failures:
            failed_lines = "\n".join(
                f"{os.path.basename(path)}: {error}" for path, error in failures
            )
            self.show_popup(
                f"Не удалось сконвертировать:\n{failed_lines}", color="#ff0000"
            )

    def check_english_words(self):
        file_path = filedialog.askopenfilename(
            title="Выберите документ",
            filetypes=[("Word Documents", "*.docx")],
        )
        if not file_path:
            return

        words_with_pos = check_english_words(file_path)

        if not words_with_pos:
            self.show_popup("Английские слова не найдены.")
            return

        popup = ctk.CTkToplevel(self, fg_color="#2f2f2f")
        popup.iconbitmap(self.icon_path)
        popup.title("")
        popup.geometry("400x400")

        tree = ttk.Treeview(
            popup, columns=("word", "paragraph"), show="headings"
        )
        tree.heading("word", text="Слово")
        tree.heading("paragraph", text="№ параграфа")
        tree.column("word", anchor="w")
        tree.column("paragraph", anchor="center")

        for word, positions in words_with_pos.items():
            paragraphs = ", ".join(str(p) for p, _ in positions)
            tree.insert("", "end", values=(word, paragraphs))

        tree.pack(expand=True, fill="both", padx=10, pady=10)

        button_frame = ctk.CTkFrame(popup, fg_color="#2f2f2f")
        button_frame.pack(pady=(0, 10))

        if len(words_with_pos) > 50:
            save_button = ctk.CTkButton(
                button_frame,
                text="Сохранить",
                command=lambda: self.save_words_to_file(words_with_pos),
                corner_radius=self.button_corner_radius,
                bg_color="#2f2f2f",
                fg_color="#313131",
                hover_color="#3e3e3e",
                text_color=self.button_text_color,
                border_width=0,
                font=self.custom_font,
                height=self.button_height,
            )
            save_button.pack(side="left", padx=(0, 10))
            self._apply_button_hover_effect(save_button)

        close_button = ctk.CTkButton(
            button_frame,
            text="Закрыть",
            command=popup.destroy,
            corner_radius=self.button_corner_radius,
            bg_color="#2f2f2f",
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        close_button.pack(side="left")
        self._apply_button_hover_effect(close_button)

    def find_formatted_separators(self):
        file_path = filedialog.askopenfilename(
            title="Выберите документ",
            filetypes=[("Word Documents", "*.docx")],
        )
        if not file_path:
            return

        document = Document(file_path)
        separators = collect_formatted_separators(document)

        if not separators:
            self.show_popup("Форматированные разделители не найдены.")
            return

        popup = ctk.CTkToplevel(self, fg_color="#2f2f2f")
        popup.iconbitmap(self.icon_path)
        popup.title("")
        popup.geometry("360x320")

        label = ctk.CTkLabel(
            popup,
            text=f"Найдено {len(separators)} форматированных разделителя(ей).",
            text_color="#eeeeee",
            font=self.custom_font,
        )
        label.pack(padx=10, pady=(10, 0))

        tree = ttk.Treeview(popup, columns=("paragraph",), show="headings")
        tree.heading("paragraph", text="№ параграфа")
        tree.column("paragraph", anchor="center")

        for index, _ in separators:
            tree.insert("", "end", values=(index,))

        tree.pack(expand=True, fill="both", padx=10, pady=10)

        button_frame = ctk.CTkFrame(popup, fg_color="#2f2f2f")
        button_frame.pack(pady=(0, 10))

        def fix():
            for _, paragraph in separators:
                fix_formatted_separator(paragraph)
            document.save(file_path)
            popup.destroy()
            self.show_popup(
                f"Устранено форматирование у {len(separators)} разделител(я/ей)."
            )

        fix_button = ctk.CTkButton(
            button_frame,
            text="Устранить форматирование",
            command=fix,
            corner_radius=self.button_corner_radius,
            bg_color="#2f2f2f",
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        fix_button.pack(side="left", padx=(0, 10))
        self._apply_button_hover_effect(fix_button)

        close_button = ctk.CTkButton(
            button_frame,
            text="Закрыть",
            command=popup.destroy,
            corner_radius=self.button_corner_radius,
            bg_color="#2f2f2f",
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        close_button.pack(side="left")
        self._apply_button_hover_effect(close_button)

    def check_duplicate_chapters(self):
        file_path = filedialog.askopenfilename(
            title="Выберите документ",
            filetypes=[("Word Documents", "*.docx")],
        )
        if not file_path:
            return

        duplicates = find_duplicate_chapters(file_path)

        if not duplicates:
            self.show_popup("Повторов не найдено.")
            return

        popup = ctk.CTkToplevel(self, fg_color="#2f2f2f")
        popup.iconbitmap(self.icon_path)
        popup.title("")
        popup.geometry("450x400")

        tree = ttk.Treeview(
            popup, columns=("chapters", "preview"), show="headings"
        )
        tree.heading("chapters", text="Главы")
        tree.heading("preview", text="Начало текста")
        tree.column("chapters", anchor="w", width=180)
        tree.column("preview", anchor="w")

        for titles, content in duplicates:
            snippet = re.sub(r"\s+", " ", content).strip()
            if len(snippet) > 120:
                snippet = snippet[:117] + "…"
            if not snippet:
                snippet = "(пусто)"
            tree.insert("", "end", values=(", ".join(titles), snippet))

        tree.pack(expand=True, fill="both", padx=10, pady=10)

        button_frame = ctk.CTkFrame(popup, fg_color="#2f2f2f")
        button_frame.pack(pady=(0, 10))

        close_button = ctk.CTkButton(
            button_frame,
            text="Закрыть",
            command=popup.destroy,
            corner_radius=self.button_corner_radius,
            bg_color="#2f2f2f",
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        close_button.pack(side="left")
        self._apply_button_hover_effect(close_button)

    def check_chapter_numbering(self):
        file_path = filedialog.askopenfilename(
            title="Выберите документ",
            filetypes=[("Word Documents", "*.docx")],
        )
        if not file_path:
            return

        missing = find_missing_chapters(file_path)

        if not missing:
            self.show_popup("Все ровно!")
            return

        message = "Отсутствуют следующие главы:\n" + "\n".join(missing)
        self.show_popup(message, color="#ff0000")

    def save_words_to_file(self, words):
        folder = filedialog.askdirectory(title="Выберите папку для сохранения списка")
        if not folder:
            return
        file_path = os.path.join(folder, "english_words.txt")
        lines = [
            f"{word} - {', '.join(str(p) for p, _ in positions)}"
            for word, positions in words.items()
        ]
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        self.show_popup(f"Список сохранен в {file_path}")

    def open_upload_dialog(self):
        files = filedialog.askopenfilenames(
            title="Выберите главы",
            filetypes=[("Word Documents", "*.docx")],
        )
        if not files:
            return

        dialog = ctk.CTkToplevel(self, fg_color="#2f2f2f")
        dialog.iconbitmap(self.icon_path)
        dialog.title("")

        # Input fields
        inputs = {}
        fields = [
            ("URL книги", "book_url", False),
            ("Логин", "username", False),
            ("Пароль", "password", True),
            ("Том", "volume", False),
            ("Дата/время публикации", "publish_at", False),
        ]
        for label_text, key, is_password in fields:
            label = ctk.CTkLabel(dialog, text=label_text, text_color="#eeeeee", font=self.custom_font)
            label.pack(padx=10, pady=(10, 0), anchor="w")
            entry = ctk.CTkEntry(
                dialog,
                show="*" if is_password else None,
                fg_color="#ffffff",
                border_color="#2f2f2f",
                text_color="#303030",
                corner_radius=self.entry_corner_radius,
                border_width=0,
                font=self.custom_font,
                height=self.entry_height,
            )
            entry.pack(fill=tk.X, padx=10, pady=(0, 10))

            # Enable clipboard operations
            entry.bind("<Control-v>", lambda e: e.widget.event_generate("<<Paste>>"))
            entry.bind("<Control-V>", lambda e: e.widget.event_generate("<<Paste>>"))

            def _show_menu(event, widget=entry):
                menu = tk.Menu(widget, tearoff=0)
                menu.add_command(
                    label="Вставить", command=lambda: widget.event_generate("<<Paste>>")
                )
                menu.tk_popup(event.x_root, event.y_root)

            entry.bind("<Button-3>", _show_menu)
            inputs[key] = entry

        deferred_var = tk.BooleanVar()
        subscription_var = tk.BooleanVar()
        deferred_cb = ctk.CTkCheckBox(
            dialog,
            text="Отложенная публикация",
            variable=deferred_var,
            text_color="#eeeeee",
            fg_color="#313131",
            hover_color="#3e3e3e",
            border_width=0,
            font=self.custom_font,
        )
        deferred_cb.pack(padx=10, pady=(0, 5), anchor="w")

        subscription_cb = ctk.CTkCheckBox(
            dialog,
            text="Подписка",
            variable=subscription_var,
            text_color="#eeeeee",
            fg_color="#313131",
            hover_color="#3e3e3e",
            border_width=0,
            font=self.custom_font,
        )
        subscription_cb.pack(padx=10, pady=(0, 10), anchor="w")

        button_frame = ctk.CTkFrame(dialog, fg_color="#2f2f2f")
        button_frame.pack(pady=(0, 10))

        def submit():
            book_url = inputs["book_url"].get().strip()
            username = inputs["username"].get().strip() or None
            password = inputs["password"].get() or None
            volume_text = inputs["volume"].get().strip()
            volume = int(volume_text) if volume_text else None
            publish_at = inputs["publish_at"].get().strip() or None
            deferred = bool(deferred_var.get())
            subscription = bool(subscription_var.get())

            dialog.destroy()

            try:
                results = upload_chapters(
                    book_url,
                    files,
                    username=username,
                    password=password,
                    deferred=deferred,
                    subscription=subscription,
                    volume=volume,
                    publish_at=publish_at,
                )
            except Exception as exc:  # pragma: no cover - external interaction
                self.show_popup(str(exc), color="#ff0000")
                return

            popup = ctk.CTkToplevel(self, fg_color="#2f2f2f")
            popup.iconbitmap(self.icon_path)
            popup.title("")
            lines = [
                f"{os.path.basename(path)}: {'успех' if ok else 'ошибка'}"
                for path, ok in results.items()
            ]
            label = ctk.CTkLabel(
                popup,
                text="\n".join(lines) or "Нет результатов",
                text_color="#eeeeee",
                justify="left",
                font=self.custom_font,
            )
            label.pack(padx=20, pady=20)
            close_btn = ctk.CTkButton(
                popup,
                text="Закрыть",
                command=popup.destroy,
                corner_radius=self.button_corner_radius,
                bg_color="#2f2f2f",
                fg_color="#313131",
                hover_color="#3e3e3e",
                text_color=self.button_text_color,
                border_width=0,
                font=self.custom_font,
                height=self.button_height,
            )
            close_btn.pack(pady=(0, 10))
            self._apply_button_hover_effect(close_btn)

        def cancel():
            dialog.destroy()

        ok_button = ctk.CTkButton(
            button_frame,
            text="OK",
            command=submit,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        ok_button.pack(side="left", padx=(0, 10))
        self._apply_button_hover_effect(ok_button)

        cancel_button = ctk.CTkButton(
            button_frame,
            text="Cancel",
            command=cancel,
            corner_radius=self.button_corner_radius,
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
            height=self.button_height,
        )
        cancel_button.pack(side="left")
        self._apply_button_hover_effect(cancel_button)

    def ask_questions(self):
        total_dialog = CustomInputDialog(
            self, "Сколько ебануть?", self.custom_font, self.icon_path
        )

        total_chapters = total_dialog.get_input()
        if total_chapters is None:
            return
        total_chapters = int(total_chapters)

        parts_dialog = CustomInputDialog(
            self,
            "На сколько частей делим?",
            self.custom_font,
            self.icon_path,
        )

        parts_per_chapter = parts_dialog.get_input()
        if parts_per_chapter is None:
            return
        parts_per_chapter = int(parts_per_chapter)

        save_location = self.path_entry.get()

        if save_location:
            self.generate_files(save_location, total_chapters, parts_per_chapter)
        else:
            self.show_error("Не выбрана папка для сохранения.")

    def generate_files(self, save_location, total_chapters, parts_per_chapter):
        if not os.path.exists(save_location):
            os.makedirs(save_location)

        # Создание подпапки для файлов
        timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")
        folder_for_chapters = os.path.join(save_location, f"Генерация_{timestamp}")
        os.makedirs(folder_for_chapters)

        for chapter in range(1, total_chapters + 1):
            for part in range(1, parts_per_chapter + 1):
                file_name = f"Глава {chapter}.{part}.docx"
                file_path = os.path.join(folder_for_chapters, file_name)

                doc = Document()
                # Пустой документ с минимальным содержимым (например, один пробел)
                doc.add_paragraph(" ")  # Добавляем один пробел
                doc.save(file_path)

        self.show_message(f"Создано {total_chapters * parts_per_chapter} файлов в папке {folder_for_chapters}")

    def show_message(self, message):
        self.show_popup(message)

    def show_error(self, message):
        self.show_popup(message, color="#ff0000")

    def on_closing(self):
        self.config_data["geometry"] = self.geometry()
        self.config_data["font_size"] = str(self.custom_font.cget("size"))
        self.save_config()
        self.destroy()

    def show_popup(self, message, color="#00ff00"):
        popup = ctk.CTkToplevel(self, fg_color="#2f2f2f")
        popup.iconbitmap(self.icon_path)
        popup.title("")
        popup.geometry("300x100")

        frame = ctk.CTkFrame(popup, corner_radius=12, fg_color="#2f2f2f")
        frame.pack(fill="both", expand=True)

        label = ctk.CTkLabel(
            frame,
            text=message,
            text_color=color,
            font=ctk.CTkFont(
                family=self.custom_font.actual("family"),
                size=self.custom_font.cget("size"),
                weight="bold",
            ),
        )
        label.pack(pady=20)

        close_button = ctk.CTkButton(
            frame,
            text="Закрыть",
            command=popup.destroy,
            corner_radius=self.button_corner_radius,
            bg_color="#2f2f2f",
            fg_color="#313131",
            hover_color="#3e3e3e",
            text_color=self.button_text_color,
            border_width=0,
            font=self.custom_font,
        )
        close_button.pack(pady=5)
        close_button.configure(height=self.button_height)
        self._apply_button_hover_effect(close_button)

    def load_config(self):
        config = {}
        if os.path.exists(CONFIG_PATH):
            with open(CONFIG_PATH) as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    if "=" in line:
                        key, value = line.split("=", 1)
                        config[key] = value
                    else:
                        config["geometry"] = line
        return config

    def save_config(self):
        with open(CONFIG_PATH, "w") as f:
            for key, value in self.config_data.items():
                f.write(f"{key}={value}\n")

if __name__ == "__main__":
    app = Application()
    app.mainloop()
